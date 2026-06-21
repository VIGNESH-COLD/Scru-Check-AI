"""
Script to generate sample test documents for ScruCheck AI
Run this to create sample Question Paper and Syllabus files
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_question_paper():
    """Create a sample question paper document."""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ANNA UNIVERSITY\n")
    run.bold = True
    run.font.size = Pt(14)
    
    header.add_run("B.E./B.Tech. DEGREE EXAMINATIONS, NOVEMBER 2024\n")
    header.add_run("Third Semester\n")
    header.add_run("Electrical and Electronics Engineering\n")
    header.add_run("GE23211 - BASIC ELECTRICAL AND ELECTRONICS ENGINEERING\n")
    header.add_run("(Regulation R-2021)\n")
    header.add_run("Time: Three Hours                Maximum: 100 Marks\n")
    header.add_run("Answer ALL Questions\n")
    
    doc.add_paragraph()
    
    # Part A
    part_a = doc.add_paragraph()
    part_a_run = part_a.add_run("PART A - (10 × 2 = 20 Marks)")
    part_a_run.bold = True
    
    questions_part_a = [
        "1. Define RMS value of an AC current.",
        "2. Calculate the equivalent resistance when two resistors of 4Ω and 6Ω are connected in parallel.",
        "3. State Faraday's laws of electromagnetic induction.",
        "4. What is the principle of a DC generator?",
        "5. List the types of DC motors based on excitation.",
        "6. Define power factor in AC circuits.",
        "7. What is the function of a transformer?",
        "8. Identify the main parts of a single-phase induction motor.",
        "9. State Kirchhoff's current law (KCL).",
        "10. What is impedance in AC circuits?"
    ]
    
    for q in questions_part_a:
        doc.add_paragraph(q)
    
    doc.add_paragraph()
    
    # Part B
    part_b = doc.add_paragraph()
    part_b_run = part_b.add_run("PART B - (5 × 16 = 80 Marks)")
    part_b_run.bold = True
    
    questions_part_b = [
        "11. (a) Explain the construction and working principle of a DC generator with a neat diagram. Derive the EMF equation. (16 Marks)",
        "12. (a) Analyze the operation of a single-phase transformer. Derive the EMF equation and explain the losses in a transformer. (16 Marks)",
        "13. (a) Describe the construction and working of a three-phase induction motor. Explain the concept of slip and its significance. (16 Marks)",
        "14. (a) Apply Kirchhoff's laws to solve the following circuit. Find the current in each branch using mesh analysis. [Refer Fig.1] (16 Marks)",
        "15. (a) Design a simple RLC series circuit and calculate the resonant frequency, bandwidth, and quality factor for given values. (16 Marks)"
    ]
    
    for q in questions_part_b:
        doc.add_paragraph(q)
    
    doc.add_paragraph()
    doc.add_paragraph("Note: Scientific calculator is permitted.")
    
    # Save
    filepath = "samples/sample_question_paper.docx"
    os.makedirs("samples", exist_ok=True)
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_sample_syllabus():
    """Create a sample syllabus document."""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("GE23211 - BASIC ELECTRICAL AND ELECTRONICS ENGINEERING\n")
    run.bold = True
    run.font.size = Pt(14)
    
    header.add_run("SYLLABUS\n")
    header.add_run("Regulation: R-2021\n")
    header.add_run("Credits: 4 | L-T-P: 3-1-0\n")
    
    doc.add_paragraph()
    
    # Course Outcomes
    co_header = doc.add_paragraph()
    co_header.add_run("COURSE OUTCOMES:").bold = True
    
    course_outcomes = [
        "CO1: Apply fundamental concepts of electrical circuits to solve problems.",
        "CO2: Understand the working principles of DC machines and transformers.",
        "CO3: Analyze the operation and characteristics of AC machines.",
        "CO4: Design and evaluate simple electronic circuits.",
        "CO5: Apply electromagnetic principles to real-world applications."
    ]
    
    for co in course_outcomes:
        doc.add_paragraph(co)
    
    doc.add_paragraph()
    
    # Unit I
    unit1 = doc.add_paragraph()
    unit1.add_run("UNIT I - ELECTRICAL CIRCUITS (9 Hours)").bold = True
    doc.add_paragraph("""
Ohm's Law - Resistors in series and parallel - Voltage and current divider rules - 
Kirchhoff's laws - Mesh and nodal analysis - Network theorems: Thevenin, Norton, 
Superposition - AC fundamentals - RMS and average values - Phasors - Impedance - 
Power factor - Power in AC circuits - Single phase and three phase systems.
    """)
    
    # Unit II
    unit2 = doc.add_paragraph()
    unit2.add_run("UNIT II - DC MACHINES (9 Hours)").bold = True
    doc.add_paragraph("""
DC Generator: Construction, principle of operation, EMF equation, types of generators.
DC Motor: Principle of operation, back EMF, torque equation, types of DC motors,
speed control methods, applications. Losses and efficiency of DC machines.
    """)
    
    # Unit III
    unit3 = doc.add_paragraph()
    unit3.add_run("UNIT III - TRANSFORMERS (9 Hours)").bold = True
    doc.add_paragraph("""
Single phase transformer: Construction, principle of operation, EMF equation,
transformation ratio. Losses in transformers - Copper loss, Iron loss.
Efficiency and voltage regulation. Open circuit and short circuit tests.
Auto-transformer and its applications.
    """)
    
    # Unit IV
    unit4 = doc.add_paragraph()
    unit4.add_run("UNIT IV - AC MACHINES (9 Hours)").bold = True
    doc.add_paragraph("""
Three-phase induction motor: Construction, principle of operation, slip, torque-slip
characteristics. Starting methods of induction motors. Single-phase induction motor:
Construction, principle, starting methods. Synchronous machines: Principle of operation,
applications.
    """)
    
    # Unit V
    unit5 = doc.add_paragraph()
    unit5.add_run("UNIT V - BASICS OF ELECTRONICS (9 Hours)").bold = True
    doc.add_paragraph("""
PN junction diode: Characteristics and applications. Zener diode.
Bipolar Junction Transistor (BJT): Construction, working, characteristics.
Field Effect Transistor (FET): Construction, working, characteristics.
Basic amplifier circuits. Operational amplifier fundamentals.
    """)
    
    # Save
    filepath = "samples/sample_syllabus.docx"
    os.makedirs("samples", exist_ok=True)
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_sample_previous_paper():
    """Create a sample previous year question paper."""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("ANNA UNIVERSITY\n")
    run.bold = True
    run.font.size = Pt(14)
    
    header.add_run("B.E./B.Tech. DEGREE EXAMINATIONS, APRIL 2024\n")
    header.add_run("Third Semester\n")
    header.add_run("GE23211 - BASIC ELECTRICAL AND ELECTRONICS ENGINEERING\n")
    header.add_run("(Previous Year Paper)\n")
    
    doc.add_paragraph()
    
    # Part A (some similar questions for repetition detection)
    part_a = doc.add_paragraph()
    part_a_run = part_a.add_run("PART A - (10 × 2 = 20 Marks)")
    part_a_run.bold = True
    
    questions_part_a = [
        "1. What is RMS value of alternating current?",  # Similar to current paper Q1
        "2. Find equivalent resistance of 3Ω and 6Ω resistors in parallel.",
        "3. State Lenz's law.",
        "4. Explain the working principle of DC generator.",  # Similar to current paper Q4
        "5. List the applications of DC series motor.",
        "6. What is power factor? Why is it important?",  # Similar to current paper Q6
        "7. Define transformation ratio of a transformer.",
        "8. What is slip in an induction motor?",
        "9. State Kirchhoff's voltage law (KVL).",
        "10. Define reactance in AC circuits."
    ]
    
    for q in questions_part_a:
        doc.add_paragraph(q)
    
    # Save
    filepath = "samples/sample_previous_paper.docx"
    os.makedirs("samples", exist_ok=True)
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


if __name__ == "__main__":
    print("Generating sample test documents...")
    create_sample_question_paper()
    create_sample_syllabus()
    create_sample_previous_paper()
    print("\nAll sample files created in 'samples/' folder!")
    print("\nYou can now upload these files in the ScruCheck AI frontend.")
