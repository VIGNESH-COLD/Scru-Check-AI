"""
Document Parser Service
Extracts text from PDF and DOCX files
"""

import io
from typing import List, Dict, Any
from fastapi import UploadFile
import fitz
import pymupdf4llm
from docx import Document


class DocumentParser:
    """Parse PDF and DOCX documents to extract text content."""
    
    async def parse(self, file: UploadFile) -> Dict[str, Any]:
        """
        Parse uploaded file and extract text content.
        
        Returns:
            Dict with filename, content, and extracted questions
        """
        content = await file.read()
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            text, image_info = self._parse_pdf(content)
        elif filename.endswith('.docx'):
            text, image_info = self._parse_docx(content)
        else:
            raise ValueError(f"Unsupported file format: {filename}")
        
        # Extract questions from text
        questions = self._extract_questions(text)
        
        # Debug logging
        print(f"📄 DocumentParser: extracted {len(questions)} questions from {file.filename}")
        print(f"📄 DocumentParser: raw_text length = {len(text)}")
        print(f"📄 DocumentParser: image_info = {image_info}")
        
        return {
            "filename": file.filename,
            "raw_text": text,
            "questions": questions,
            "sections": self._identify_sections(text),
            "image_info": image_info
        }
    
    def _parse_pdf(self, content: bytes):
        """Extract text and detect embedded images from PDF file using pymupdf4llm."""
        doc = fitz.open(stream=content, filetype="pdf")
        
        # Use pymupdf4llm to extract text as Markdown
        text = pymupdf4llm.to_markdown(doc)
        
        image_count = 0
        broken_images = 0

        for page in doc:
            try:
                images = page.get_images(full=True)
                image_count += len(images)
            except Exception:
                broken_images += 1
        
        image_info = {
            "has_images": image_count > 0 or broken_images > 0,
            "image_count": image_count,
            "broken_images": broken_images,
            "source": "pymupdf"
        }
        return text, image_info
    
    def _parse_docx(self, content: bytes):
        """Extract text and detect embedded images from DOCX file."""
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        # Detect inline images via inline_shapes
        image_count = 0
        broken_images = 0
        try:
            for shape in doc.inline_shapes:
                # WD_INLINE_SHAPE_TYPE.PICTURE == 3
                if shape.type is not None and int(shape.type) == 3:
                    # Validate size > 0 means the image is properly embedded
                    w = shape.width or 0
                    h = shape.height or 0
                    if w > 0 and h > 0:
                        image_count += 1
                    else:
                        broken_images += 1
        except Exception:
            pass  # inline_shapes not available or error

        image_info = {
            "has_images": image_count > 0 or broken_images > 0,
            "image_count": image_count,
            "broken_images": broken_images,
            "source": "docx_inline_shapes"
        }
        return "\n".join(text_parts), image_info
    
    def _extract_questions(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract individual questions from document text.
        Handles various numbering formats: 1., Q1, 1), (a), etc.
        """
        import re
        
        # Pre-process text to ensure questions on the same line are split onto new lines
        # This handles cases where PDF extraction merges everything onto one line
        text = re.sub(r'(\s+|^)(\d+\([a-zA-Z]\)\.?|\d+\.|\d+\)|Q\d+\.)\s+', r'\n\2 ', text)
        
        # Force section headers and metadata to start on new lines so they can be filtered
        text = re.sub(r'(?i)(\s+)(PART\s+[A-Z]|SECTION\s+[A-Z]|Time:|Maximum|Max Marks|Note:|Answer\b)', r'\n\2', text)
        
        questions = []
        lines = text.split('\n')
        
        # Patterns for question detection
        # Patterns for question detection (allow leading whitespace)
        patterns = [
            r'^\s*(\d+\([a-zA-Z]\)\.?)\s+(.+)', # 11(a). or 11(a) Question text
            r'^\s*(\d+)\.\s+(.+)',           # 1. Question text
            r'^\s*Q(\d+)\.\s*(.+)',           # Q1. Question text
            r'^\s*(\d+)\)\s+(.+)',            # 1) Question text
            r'^\s*\(([a-z])\)\s+(.+)',        # (a) Question text
            r'^\s*([ivx]+)\.\s+(.+)',         # i. Question text (roman numerals)
        ]
        
        current_question = None
        question_num = 0
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            matched = False
            for pattern in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    if current_question:
                        questions.append(current_question)
                    
                    question_num += 1
                    current_question = {
                        "number": question_num,
                        "label": match.group(1),
                        "text": match.group(2),
                        "full_text": line
                    }
                    matched = True
                    break
            
            # Continue building current question if it spans multiple lines
            if not matched and current_question:
                # Filter out section headers and common metadata
                if re.match(r'^\s*[\(\[]?(PART|SECTION)\s+[A-Z]\b', line, re.IGNORECASE):
                    continue
                if re.match(r'^\s*[\(\[]?(Time:|Maximum|Max Marks|Note:|Answer\b)', line, re.IGNORECASE):
                    continue
                current_question["text"] += " " + line
                current_question["full_text"] += " " + line
        
        if current_question:
            questions.append(current_question)
            
        print(f"📄 Document Parser: Extracted {len(questions)} questions from {len(lines)} lines")
        if len(questions) == 0:
            print(f"❌ raw_text snippet: {text[:500]}")
        return questions
    
    def _identify_sections(self, text: str) -> List[Dict[str, str]]:
        """Identify section headers like Part A, Part B, Section A, etc."""
        import re
        
        sections = []
        patterns = [
            r'(PART\s+[A-Z])',
            r'(SECTION\s+[A-Z])',
            r'(Part\s+[A-Z])',
            r'(Section\s+[A-Z])',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if match.upper() not in [s["name"].upper() for s in sections]:
                    sections.append({"name": match.strip()})
        
        return sections
