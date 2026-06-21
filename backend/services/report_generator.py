"""
Report Generator
Generates DOCX scrutiny reports with two-section layout:
  Section 1: Mandatory Compliance (PASS/FAIL)
  Section 2: Quality Scores (0-100)
"""

from typing import Dict, Any, List
from datetime import datetime
import os
import matplotlib.pyplot as plt
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class ReportGenerator:
    """Generate compliance reports in DOCX format with charts."""
    
    REPORTS_DIR = "reports"
    TEMP_IMG_DIR = "temp_images"
    
    def __init__(self):
        os.makedirs(self.REPORTS_DIR, exist_ok=True)
        os.makedirs(self.TEMP_IMG_DIR, exist_ok=True)
    
    async def generate(self, paper_id: str, format: str = "docx", data: Dict[str, Any] = None) -> str:
        """Generate report and return file path."""
        
        doc = Document()
        
        # Default mock data if not provided (fallback)
        if not data:
            data = {
                "criteria": [],
                "mandatory_compliance": [],
                "quality_scores": [],
                "blooms": {'Remember': 0, 'Understand': 0, 'Apply': 0, 'Analyze': 0, 'Evaluate': 0, 'Create': 0},
                "syllabus_coverage": {},
                "score": "N/A"
            }

        # --- Titles & Header ---
        title = doc.add_heading("ScruCheck AI - Scrutiny Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Paper ID: {paper_id}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        overall_status = data.get("overall_status", "N/A")
        status_emoji = {"APPROVED": "✅", "CONDITIONAL": "⚠️", "REJECTED": "❌"}.get(overall_status, "")
        doc.add_paragraph(f"Status: {overall_status} {status_emoji}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        score_str = data.get("score", "N/A")
        doc.add_paragraph(f"Score: {score_str}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # ═══════════════════════════════════════════════
        # Section 1: Mandatory Compliance
        # ═══════════════════════════════════════════════
        doc.add_heading("Section 1: Mandatory Compliance Status", level=1)
        
        mandatory_items = data.get("mandatory_compliance", [])
        if not mandatory_items:
            # Fallback: extract from criteria
            mandatory_keys = {"format_compliance", "regulation_check", "mark_distribution", "permitted_aids"}
            mandatory_items = [
                {"criterion": c["criterion"], "status": c["status"], "remarks": c.get("remarks", c.get("suggestion", ""))}
                for c in data.get("criteria", []) if c["criterion"] in mandatory_keys
            ]
        
        mandatory_labels = {
            "format_compliance": "Format Compliance",
            "regulation_check": "Regulation Compliance",
            "mark_distribution": "Mark Distribution",
            "permitted_aids": "Permitted Aids Compliance",
        }
        
        num_mandatory = max(len(mandatory_items), 4)
        table1 = doc.add_table(rows=num_mandatory + 1, cols=3)
        table1.style = 'Table Grid'
        
        # Headers
        h1 = table1.rows[0].cells
        h1[0].text = "Criterion"
        h1[1].text = "Status"
        h1[2].text = "Remarks"
        for cell in h1:
            cell.paragraphs[0].runs[0].bold = True
        
        for i, item in enumerate(mandatory_items, 1):
            if i >= len(table1.rows):
                break
            row = table1.rows[i].cells
            row[0].text = mandatory_labels.get(item["criterion"], item["criterion"])
            status = item.get("status", "N/A")
            if status == "PASS":
                row[1].text = "✅ PASS"
            elif status == "FAIL":
                row[1].text = "❌ FAIL"
            else:
                row[1].text = status
            row[2].text = item.get("remarks", "")
        
        doc.add_paragraph()
        
        # ═══════════════════════════════════════════════
        # Section 2: Quality Scores
        # ═══════════════════════════════════════════════
        doc.add_heading("Section 2: Quality Scores", level=1)
        
        quality_items = data.get("quality_scores", [])
        if not quality_items:
            # Fallback: extract from criteria
            quality_keys = {"syllabus_alignment", "blooms_taxonomy", "grammar_clarity", "repetition_check", "diagrams_symbols"}
            quality_items = [
                {"criterion": c["criterion"], "score": c.get("score", 0), "remarks": c.get("remarks", c.get("suggestion", ""))}
                for c in data.get("criteria", []) if c["criterion"] in quality_keys
            ]
        
        quality_labels = {
            "syllabus_alignment": "Syllabus Coverage Score",
            "blooms_taxonomy": "Bloom's Taxonomy Distribution Score",
            "grammar_clarity": "Grammar and Clarity Score",
            "repetition_check": "Repetition Risk Score",
            "diagrams_symbols": "Diagram and Symbol Quality Score",
        }
        
        num_quality = max(len(quality_items), 5)
        table2 = doc.add_table(rows=num_quality + 1, cols=3)
        table2.style = 'Table Grid'
        
        h2 = table2.rows[0].cells
        h2[0].text = "Criterion"
        h2[1].text = "Score"
        h2[2].text = "Remarks"
        for cell in h2:
            cell.paragraphs[0].runs[0].bold = True
        
        for i, item in enumerate(quality_items, 1):
            if i >= len(table2.rows):
                break
            row = table2.rows[i].cells
            row[0].text = quality_labels.get(item["criterion"], item["criterion"])
            score_val = item.get("score", 0)
            row[1].text = f"{score_val}/100"
            row[2].text = item.get("remarks", "")
        
        doc.add_paragraph()

        # --- Bloom's Taxonomy Visualization ---
        doc.add_heading("Cognitive Level Analysis (Bloom's Taxonomy)", level=1)
        
        # Generate Chart
        blooms_data = data.get("blooms", {})
        if not blooms_data:
             blooms_data = {'Remember': 0, 'Understand': 0, 'Apply': 0, 'Analyze': 0, 'Evaluate': 0, 'Create': 0}
             
        blooms_img_path = self._create_blooms_chart(blooms_data, paper_id)
        
        # Add to DOCX
        doc.add_picture(blooms_img_path, width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 1: Distribution of marks across cognitive levels.")
        
        # --- Syllabus Coverage Visualization ---
        doc.add_heading("Syllabus Coverage", level=1)
        
        # Generate Chart
        syllabus_data = data.get("syllabus_coverage", {})
        if not syllabus_data:
            syllabus_data = {"Unit I": 0, "Unit II": 0}

        syllabus_img_path = self._create_syllabus_chart(syllabus_data, paper_id)
        
        # Add to DOCX
        doc.add_picture(syllabus_img_path, width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 2: Weightage of marks across syllabus units.")
        
        doc.add_paragraph()

        # --- Recommendations ---
        doc.add_heading("AI Recommendations", level=1)
        
        has_recommendations = False
        
        # Mandatory failures
        for item in mandatory_items:
            if item.get("status") == "FAIL":
                has_recommendations = True
                doc.add_paragraph(f"• ❌ {mandatory_labels.get(item['criterion'], item['criterion'])}: {item.get('remarks', 'Manual review required.')}")
        
        # Syllabus floor rules (independent of avg quality)
        syllabus_score = next(
            (item.get("score") for item in quality_items if item.get("criterion") == "syllabus_alignment"),
            None
        )
        if syllabus_score is not None and syllabus_score < 40:
            has_recommendations = True
            doc.add_paragraph(
                f"• 🚨 CRITICAL: Syllabus Coverage is {syllabus_score}/100 (below 40 threshold). "
                "This paper covers fewer than 40% of the prescribed syllabus topics and has been REJECTED."
            )
        elif syllabus_score is not None and syllabus_score < 50:
            has_recommendations = True
            doc.add_paragraph(
                f"• ⚠️ Syllabus Coverage is {syllabus_score}/100 (below 50 threshold). "
                "Insufficient syllabus coverage has triggered a CONDITIONAL status."
            )
        
        # Quality scores below 70
        for item in quality_items:
            score_val = item.get("score", 0)
            if score_val < 70:
                has_recommendations = True
                label = quality_labels.get(item["criterion"], item["criterion"])
                doc.add_paragraph(f"• ⚠️ {label} ({score_val}/100): {item.get('remarks', 'Needs improvement.')}")
        
        # Review items from criteria
        review_items = [c for c in data.get("criteria", []) if c.get("status") in ["UNCERTAIN", "NOT_EVALUATED"]]
        for ri in review_items:
            has_recommendations = True
            sugg = ri.get("suggestion", "Manual review recommended.")
            doc.add_paragraph(f"• 🔍 {ri['criterion']} (Needs Review): {sugg}")
        
        if not has_recommendations:
            doc.add_paragraph("• No critical issues found. Paper is well-structured.")
        
        doc.add_paragraph()
        doc.add_paragraph("---")
        doc.add_paragraph("Generated by ScruCheck AI v1.0")

        # Save Report
        filepath = os.path.join(self.REPORTS_DIR, f"ScruCheck_Report_{paper_id}.docx")
        doc.save(filepath)
        
        # Cleanup temp images
        try:
            os.remove(blooms_img_path)
            os.remove(syllabus_img_path)
        except:
            pass
            
        return filepath

    def _create_blooms_chart(self, data: Dict[str, int], paper_id: str) -> str:
        """Create a pie chart for Bloom's Taxonomy."""
        labels = list(data.keys())
        sizes = list(data.values())
        colors = ['#8884d8', '#82ca9d', '#ffc658', '#ff8042', '#0088fe', '#00C49F']
        
        plt.figure(figsize=(6, 4))
        plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
        plt.axis('equal')
        plt.title("Bloom's Taxonomy Distribution")
        
        filename = f"blooms_{paper_id}.png"
        path = os.path.join(self.TEMP_IMG_DIR, filename)
        plt.savefig(path, bbox_inches='tight')
        plt.close()
        return path

    def _create_syllabus_chart(self, data: Dict[str, int], paper_id: str) -> str:
        """Create a bar chart for Syllabus Coverage."""
        units = list(data.keys())
        marks = list(data.values())
        
        plt.figure(figsize=(6, 4))
        plt.bar(units, marks, color='#3b82f6')
        plt.xlabel('Units')
        plt.ylabel('Marks')
        plt.title('Syllabus Mark Weightage')
        
        filename = f"syllabus_{paper_id}.png"
        path = os.path.join(self.TEMP_IMG_DIR, filename)
        plt.savefig(path, bbox_inches='tight')
        plt.close()
        return path

# Instance for main.py to import
report_generator = ReportGenerator()
