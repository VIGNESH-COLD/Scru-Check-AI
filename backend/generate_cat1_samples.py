import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cat1_samples():
    output_dir = os.path.join(os.path.dirname(__file__), "samples")
    os.makedirs(output_dir, exist_ok=True)

    # --- 1. Generate Syllabus (Unit I & II) ---
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    heading = doc.add_heading('SYLLABUS - PYTHON PROGRAMMING', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Unit I
    doc.add_heading('UNIT I - ALGORITHMIC PROBLEM SOLVING', 1)
    doc.add_paragraph('Algorithms, building blocks of algorithms (statements, state, control flow, functions), notation (pseudo code, flow chart, programming language), algorithmic problem solving, simple strategies for developing algorithms (iteration, recursion). Illustrative problems: find minimum in a list, insert a card in a list of sorted cards, guess an integer number in a range, Towers of Hanoi.')

    # Unit II
    doc.add_heading('UNIT II - DATA, EXPRESSIONS, STATEMENTS', 1)
    doc.add_paragraph('Python interpreter and interactive mode; values and types: int, float, boolean, string, and list; variables, expressions, statements, tuple assignment, precedence of operators, comments; modules and functions, function definition and use, flow of execution, parameters and arguments; Illustrative programs: exchange the values of two variables, circulate the values of n variables, distance between two points.')

    doc.save(os.path.join(output_dir, "CAT1_Syllabus.docx"))
    print(f"Created: {os.path.join(output_dir, 'CAT1_Syllabus.docx')}")

    # --- 2. Generate Question Paper (CAT 1 Standard) ---
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header
    head = doc.add_heading('CAT-1 EXAMINATION', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Course: Python Programming\nMax Marks: 50\nTime: 90 Minutes')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PART A (Unit I)
    doc.add_heading('PART A (10 Marks) - From Unit I', 1)
    doc.add_paragraph('(Answer ALL questions. Each carries 2 marks)')
    
    questions_u1_part_a = [
        "1. Define an algorithm.",
        "2. What is a flowchart? Draw the symbol for decision making.",
        "3. Differentiate between pseudocode and code.",
        "4. List the building blocks of algorithms.",
        "5. What is recursion? Give an example."
    ]
    for q in questions_u1_part_a:
        doc.add_paragraph(q)

    # PART B (Unit I)
    doc.add_heading('PART B (15 Marks) - From Unit I', 1)
    doc.add_paragraph('(Answer Any ONE Question. Carries 15 marks)')
    
    questions_u1_part_b = [
        "6. (a) Explain algorithmic problem solving strategies with examples. (15)",
        "OR",
        "6. (b) Write an algorithm and draw a flowchart to find the minimum in a list. (15)"
    ]
    for q in questions_u1_part_b:
        doc.add_paragraph(q)

    # PART A (Unit II)
    doc.add_heading('PART A (10 Marks) - From Unit II', 1)
    doc.add_paragraph('(Answer ALL questions. Each carries 2 marks)')
    
    questions_u2_part_a = [
        "7. Define a variable in Python.",
        "8. What are the data types available in Python?",
        "9. Explain tuple assignment with an example.",
        "10. What is a function? Syntax for function definition.",
        "11. Mention the precedence of operators in Python."
    ]
    for q in questions_u2_part_a:
        doc.add_paragraph(q)

    # PART B (Unit II)
    doc.add_heading('PART B (15 Marks) - From Unit II', 1)
    doc.add_paragraph('(Answer Any ONE Question. Carries 15 marks)')
    
    questions_u2_part_b = [
        "12. (a) Explain the various operators in Python with suitable examples. (15)",
        "OR",
        "12. (b) Illustrate the concept of function with a program to circulate values of n variables. (15)"
    ]
    for q in questions_u2_part_b:
        doc.add_paragraph(q)

    doc.save(os.path.join(output_dir, "CAT1_Question_Paper.docx"))
    print(f"Created: {os.path.join(output_dir, 'CAT1_Question_Paper.docx')}")


    # --- 3. Generate Previous Question Paper (For Repetition Check) ---
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header
    head = doc.add_heading('PREVIOUS YEAR CAT-1', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('PART A', 1)
    prev_qs_a = [
        "1. What is an algorithm?", # Similar to Q1
        "2. Define recursion.", # Similar to Q5
        "3. Distinction between integer and float.",
        "4. What are keywords?",
        "5. Define function."
    ]
    for q in prev_qs_a:
        doc.add_paragraph(q)

    doc.add_heading('PART B', 1)
    prev_qs_b = [
        "6. Explain the building blocks of algorithms in detail. (15)",
        "7. Discuss the various operators available in Python. (15)" # Same as Q12a
    ]
    for q in prev_qs_b:
        doc.add_paragraph(q)

    doc.save(os.path.join(output_dir, "CAT1_Previous_Paper.docx"))
    print(f"Created: {os.path.join(output_dir, 'CAT1_Previous_Paper.docx')}")

if __name__ == "__main__":
    create_cat1_samples()
